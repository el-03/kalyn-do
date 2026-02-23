# kalyn/services/delivery_order_service.py

import os
from datetime import datetime
from pathlib import Path
from typing import List, Dict

from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv

from domain.models import DeliveryOrder, DeliveryOrderLine, BarcodeResult
from google_client import get_drive_service, get_docs_service
from services.doc_service import generate_delivery_order_doc, generate_barcode_do_doc
from utils.formatting import format_rupiah
from utils.barcode import ensure_barcode_image
from utils.docx_helpers import replace_placeholders_in_document

from services.barcode_service import get_barcode_url_list

# Resolve template paths relative to this file
BASE_DIR = Path(__file__).resolve().parent.parent
DO_TEMPLATE_PATH = BASE_DIR / "templates" / "kalyn_barcode_template.docx"
BARCODE_TEMPLATE_PATH = BASE_DIR / "templates" / "kalyn_do_template.docx"

MAX_DO_LINES = 15  # number of rows your DO template supports
MAX_BARCODE_SLOTS = 280  # number of slots your barcode template supports


def build_delivery_order_from_rows(
        outlet_name: str,
        order_rows: List[dict],
) -> DeliveryOrder:
    """
    Build a DeliveryOrder domain object from a list of UI order rows.

    Each order row is expected to look like:
      {
        "SKU": str,
        "Size": str,
        "Item": str,
        "Category": str,
        "Color": str,
        "Quantity": int,
        "Unit Price": int,
        "Total": int,
      }
    """
    lines: List[DeliveryOrderLine] = []
    grand_total = 0

    for idx, row in enumerate(order_rows, start=1):
        qty = int(row["Quantity"])
        if qty <= 0:
            continue

        unit_price = int(row["Unit Price"] or 0)
        line_total = unit_price * qty
        grand_total += line_total

        label = f"T005-{row['Category']}-{row['Item']}"

        line = DeliveryOrderLine(
            index=idx,
            label=label,
            sku=row["SKU"],
            color=row["Color"],
            size=row["Size"] or "-",
            qty=qty,
            unit_price=unit_price,
            unit_price_display=format_rupiah(unit_price),
            line_total=line_total,
            line_total_display=format_rupiah(line_total),
        )
        lines.append(line)

    return DeliveryOrder(
        outlet_name=outlet_name,
        lines=lines,
        grand_total=grand_total,
        barcodes=[]
    )


# ---------------------------------------------------------------------------
# Delivery Order DOCX
# ---------------------------------------------------------------------------

def _create_delivery_order_doc(
        delivery_order: DeliveryOrder,
        output_path: str,
) -> str:
    """
    Fill the DO Word template using the DeliveryOrder object.
    """
    doc = Document(str(DO_TEMPLATE_PATH))

    now = datetime.now()
    date_display = now.strftime("%d/%m/%Y")

    mapping: Dict[str, str] = {
        "{{date}}": date_display,
        "{{store_location}}": delivery_order.outlet_name,
        "{{total_sum}}": format_rupiah(delivery_order.grand_total),
    }

    effective_lines = delivery_order.lines[:MAX_DO_LINES]

    # Fill placeholders row by row
    for line in effective_lines:
        i = line.index  # 1-based index used in template
        mapping[f"{{{{no_{i}}}}}"] = str(i)
        mapping[f"{{{{cat_{i}}}}}"] = line.label
        mapping[f"{{{{color_{i}}}}}"] = line.color
        mapping[f"{{{{size_{i}}}}}"] = line.size
        mapping[f"{{{{harga_{i}}}}}"] = line.unit_price_display
        mapping[f"{{{{qty_{i}}}}}"] = str(line.qty)
        mapping[f"{{{{harga_sum_{i}}}}}"] = line.line_total_display
        # temporary text for barcode; will be replaced by image
        mapping[f"{{{{barcode_{i}}}}}"] = line.sku

    # Blank out unused rows
    max_used_index = max((line.index for line in effective_lines), default=0)
    for i in range(max_used_index + 1, MAX_DO_LINES + 1):
        for key_prefix in ["no", "cat", "barcode", "color", "size", "harga", "qty", "harga_sum"]:
            mapping[f"{{{{{key_prefix}_{i}}}}}"] = ""

    # First pass: replace text
    replace_placeholders_in_document(doc, mapping)

    # Second pass: insert barcode images
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text.startswith("{{barcode_") and text.endswith("}}"):
                    inside = text[2:-2]  # e.g. "barcode_3"
                    _, idx_str = inside.split("_")
                    idx = int(idx_str)

                    # clear text regardless
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = ""

                    # find line by index
                    line = next((ln for ln in effective_lines if ln.index == idx), None)
                    if not line:
                        continue

                    img_path = ensure_barcode_image(line.sku)
                    p = cell.paragraphs[0]
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(1.5))

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Barcode DOCX
# ---------------------------------------------------------------------------

def _expand_lines_for_barcode(delivery_order: DeliveryOrder) -> List[dict]:
    """
    Expand each DeliveryOrderLine into repeated label/barcode/price entries,
    one per unit (qty).
    """
    result = []
    for line in delivery_order.lines:
        for _ in range(line.qty):
            result.append(
                {
                    "label": line.label,
                    "sku": line.sku,
                    "price_display": line.unit_price_display,
                }
            )
    return result


def _create_barcode_doc(
        delivery_order: DeliveryOrder,
        output_path: str,
) -> str:
    """
    Fill the Barcode Word template using the DeliveryOrder data.
    Template is expected to have placeholders:
      {{cat_1}}, {{barcode_1}}, {{price_1}}, ... up to MAX_BARCODE_SLOTS
    """
    doc = Document(str(BARCODE_TEMPLATE_PATH))

    expanded = _expand_lines_for_barcode(delivery_order)
    n = min(MAX_BARCODE_SLOTS, len(expanded))

    mapping: Dict[str, str] = {}

    for i in range(1, n + 1):
        entry = expanded[i - 1]
        mapping[f"{{{{cat_{i}}}}}"] = entry["label"]
        mapping[f"{{{{price_{i}}}}}"] = f"Rp {entry['price_display']}"
        mapping[f"{{{{barcode_{i}}}}}"] = entry["sku"]

    # blank out remaining slots
    for i in range(n + 1, MAX_BARCODE_SLOTS + 1):
        for prefix in ["cat", "price", "barcode"]:
            mapping[f"{{{{{prefix}_{i}}}}}"] = ""

    replace_placeholders_in_document(doc, mapping)

    # Replace barcode placeholders with images
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text.startswith("{{barcode_") and text.endswith("}}"):
                    inside = text[2:-2]  # "barcode_17"
                    _, idx_str = inside.split("_")
                    idx = int(idx_str)

                    # clear existing text
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.text = ""

                    if 1 <= idx <= len(expanded):
                        sku = expanded[idx - 1]["sku"]
                        img_path = ensure_barcode_image(sku)
                        p = cell.paragraphs[0]
                        run = p.add_run()
                        run.add_picture(img_path, width=Inches(1.2))

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_documents_for_delivery_order(
        delivery_order: DeliveryOrder,
) -> dict[str, str]:
    """
    Generate both:
      - Delivery Order docx
      - Barcode labels docx

    Returns:
      {
        "do_path": "/abs/path/to/Surat Jalan ... .docx",
        "barcode_path": "/abs/path/to/Barcode ... .docx",
      }
    """
    load_dotenv()
    template_do_doc_id = os.getenv("TEMPLATE_DO_DOC_ID")
    template_barcode_doc_id = os.getenv("TEMPLATE_BARCODE_DOC_ID")

    docs_url = {
        "do_path": "",
        "barcode_path": "",
    }

    outlet_dir = {
        "Banda": ['1q-APtN37mwMl7O3IVWMKurEI7ZL_2i8S', '1-uSGu3KXngnPGWMpc1Ij4exVQ2Z4k6Mz'],
        "Karawang": ['1DUdTQKvuv4GRtjojnmHW_3BQoyyA_DCx', '1ZkTX_Uf0LgFufD4X2kz2pWWygMpHjKfj'],
        "Purwakarta": ['14ssRFy2VeBGqWJk6ahayMLX5GGhGwApM', '14TNK9i9u423-X2LC3hu6bKz7wBDhT0nF'],
    }

    drive = get_drive_service()
    docs = get_docs_service()

    do_doc_dir = outlet_dir[delivery_order.outlet_name][0]
    barcode_doc_dir = outlet_dir[delivery_order.outlet_name][1]
    delivery_order.barcodes = get_barcode_url_list(delivery_order, drive)
    docs_url["do_path"] = generate_delivery_order_doc(docs, drive, template_do_doc_id, do_doc_dir, delivery_order)
    docs_url["barcode_path"] = generate_barcode_do_doc(docs, drive, template_barcode_doc_id, barcode_doc_dir, delivery_order)

    return docs_url

    # os.makedirs(output_dir, exist_ok=True)
    #
    # timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    # safe_outlet = delivery_order.outlet_name.replace(" ", "_")
    #
    # do_filename = f"Surat Jalan - Kalyn - {safe_outlet} - {timestamp}.docx"
    # barcode_filename = f"Barcode - Kalyn - {safe_outlet} - {timestamp}.docx"
    #
    # do_path = os.path.join(output_dir, do_filename)
    # barcode_path = os.path.join(output_dir, barcode_filename)
    #
    # _create_delivery_order_doc(delivery_order, do_path)
    # _create_barcode_doc(delivery_order, barcode_path)
    #
    # return {
    #     "do_path": os.path.abspath(do_path),
    #     "barcode_path": os.path.abspath(barcode_path),
    # }
