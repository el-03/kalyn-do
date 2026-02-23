from docx import Document


def replace_placeholders_in_document(doc: Document, mapping: dict) -> None:
    """
    Replace all occurrences of keys in `mapping` with their values
    across paragraphs and tables in a python-docx Document.
    """
    # Replace in paragraphs
    for p in doc.paragraphs:
        for key, val in mapping.items():
            if key in p.text:
                for run in p.runs:
                    run.text = run.text.replace(key, val)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in mapping.items():
                    if key in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.text = run.text.replace(key, val)